from docx import Document
from docx.shared import Pt
import os

def create_docx():
    doc = Document()
    doc.add_heading('FNRC Training Data - English and Arabic', 0)

    sections = [
        {
            "heading": "ABOUT US",
            "items": [
                {
                    "title": "About The Foundation",
                    "EN": "The Fujairah Natural Resources Corporation (FNRC) was established as an independent entity by Emiri Decree No. 1 of 2009 by His Highness Sheikh Hamad bin Mohammed Al Sharqi, Ruler of Fujairah. It aims to conduct studies, research, and geological surveys to develop the rock and mineral sectors and exploit the natural resources within the Emirate's boundaries. The Foundation consists of a Council and an Executive Body.",
                    "AR": "تأسست مؤسسة الفجيرة للموارد الطبيعية كجهة مستقلة بموجب المرسوم الأميري رقم (1) لعام 2009 الصادر عن صاحب السمو الشيخ حمد بن محمد الشرقي عضو المجلس الأعلى حاكم الفجيرة. تهدف المؤسسة إلى إجراء الدراسات والبحوث والمسوحات الجيولوجية لتطوير قطاع الصخور والمعادن واستغلال الموارد الطبيعية الواقعة ضمن حدود أراضي الإمارة."
                },
                {
                    "title": "Vision and Mission",
                    "EN": "Vision: Distinctive sustainable investment for our natural resources. Mission: Regulate and invest in natural resources through excellence in services, geological research, and advanced laboratories to achieve sustainable development.",
                    "AR": "الرؤية: استثمار مستدام ومتميز لمواردنا الطبيعية. المهمة: تنظيم واستثمار الموارد الطبيعية من خلال تقديم خدمات متميزة، وبحوث جيولوجية، ومختبرات متطورة لتحقيق التنمية المستدامة."
                },
                {
                    "title": "Strategic Objectives",
                    "EN": "1. Develop the geological and mineral sector. 2. Ensure sustainable exploitation of natural resources. 3. Enhance operational efficiency and corporate excellence. 4. Support innovation and digital transformation.",
                    "AR": "الأهداف الاستراتيجية: 1. تطوير القطاع الجيولوجي والمعدني. 2. ضمان الاستغلال المستدام للموارد الطبيعية. 3. تعزيز الكفاءة التشغيلية والتميز المؤسسي. 4. دعم الابتكار والتحول الرقمي."
                }
            ]
        },
        {
            "heading": "SERVICES",
            "items": [
                {
                    "title": "Individual Services",
                    "EN": "**Request to transport materials inside and outside the Emirate (Public Transport)**: Fees range from 40 to 5000 AED depending on the request. Service takes 5 to 7 minutes.",
                    "AR": "**طلب نقل مواد داخل وخارج الإمارة (نقليات عامة)**: الرسوم تتراوح بين 40 إلى 5000 درهم حسب الطلب. تستغرق الخدمة من 5 إلى 7 دقائق."
                },
                {
                    "title": "Company Services",
                    "EN": "- No Objection Certificate (NOC): For various engineering and construction activities.\n- Blasting Operations Permit: Authorization for controlled explosions in mining or construction areas.\n- Demolition Companies Permit: Required for companies engaged in demolition within the emirate.",
                    "AR": "- شهادة لا مانع: لمختلف الأنشطة الهندسية والإنشائية.\n- تصريح عمليات التفجير: ترخيص لإجراء عمليات التفجير المنظمة في مناطق التعدين أو الإنشاءات.\n- تصريح شركات الهدم: مطلوب للشركات التي تعمل في مجال الهدم داخل الإمارة."
                }
            ]
        },
        {
            "heading": "CONTACT INFORMATION",
            "items": [
                {
                    "title": "Headquarters and Centers",
                    "EN": "- Headquarters: Fujairah, UAE.\n- Working Hours: Monday to Thursday (7:30 AM - 3:30 PM), Friday (7:30 AM - 12:00 PM).\n- Phone: +971 9 222 8333 (General Inquiries).\n- Email: info@fnrc.gov.ae\n- Customer Centers: THOBAN, SIJI, HAB-HAB, BARAK, and Dibba Center.",
                    "AR": "- المقر الرئيسي: الفجيرة، الإمارات العربية المتحدة.\n- ساعات العمل: من الاثنين إلى الخميس (7:30 صباحاً - 3:30 مساءً)، الجمعة (7:30 صباحاً - 12:00 ظهراً).\n- الهاتف: 97192228333+ (للاستفسارات العامة).\n- البريد الإلكتروني: info@fnrc.gov.ae\n- مراكز الخدمة: ثوبان، السيجي، حبحب، الحلاة، ومركز دبا."
                }
            ]
        },
        {
            "heading": "RULES AND REGULATIONS",
            "items": [
                {
                    "title": "Laws",
                    "EN": "The foundation operates under Fujairah's mining laws and environmental regulations, including decrees ensuring the protection of the environment while optimizing resource extraction.",
                    "AR": "تعمل المؤسسة بموجب قوانين التعدين واللوائح البيئية في الفجيرة، بما في ذلك المراسيم التي تضمن حماية البيئة مع تحسين استخراج الموارد."
                }
            ]
        },
        {
            "heading": "RECENT NEWS (Early 2026)",
            "items": [
                {
                    "title": "News Highlights",
                    "EN": "- 16-Mar-2026: 'Saqya Al Khair' initiative launched for community welfare.\n- 22-Feb-2026: 'Our Resources: Art and Knowledge' event combines creativity and learning.\n- 18-Feb-2026: Solidarity values promoted through joint community programs.",
                    "AR": "- 16 مارس 2026: إطلاق مبادرة 'سقيا الخير' لتعزيز العمل المجتمعي.\n- 22 فبراير 2026: فعاليات 'مواردنا فن ومعرفة' لدمج الإبداع بالتعلم.\n- 18 فبراير 2026: تعزيز قيم التكافل من خلال برامج مجتمعية مشتركة."
                }
            ]
        }
    ]

    for section in sections:
        doc.add_heading(section["heading"], level=1)
        for item in section["items"]:
            doc.add_heading(item["title"], level=2)
            
            p_en = doc.add_paragraph()
            p_en.add_run("English Context:").bold = True
            doc.add_paragraph(item["EN"])
            
            p_ar = doc.add_paragraph()
            p_ar.add_run("Arabic Context (سياق باللغة العربية):").bold = True
            doc.add_paragraph(item["AR"])

    save_path = "e:\\FNRC_Web_Projects\\DigitalHub\\FNRC_Training_Data.docx"
    doc.save(save_path)
    print(f"Created {save_path}")

if __name__ == "__main__":
    create_docx()
